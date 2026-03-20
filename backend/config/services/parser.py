# services/parser.py

import PyPDF2
import docx

# 🔥 NEW (for real PDF highlighting)
import pdfplumber


# -----------------------------
# 📄 PDF TEXT PARSER
# -----------------------------
def parse_pdf(path):

    text = ""

    try:
        with open(path, "rb") as file:

            reader = PyPDF2.PdfReader(file)

            for page in reader.pages:

                try:
                    page_text = page.extract_text()

                    if page_text:
                        text += page_text + "\n"

                except Exception as e:
                    print("⚠️ PDF page read error:", e)
                    continue

    except Exception as e:
        print("❌ PDF parsing failed:", e)
        return ""

    print("PDF TEXT LENGTH:", len(text))

    return text.strip()


# -----------------------------
# 🔥 PDF POSITION PARSER (NEW)
# Used for exact highlighting
# -----------------------------
def parse_pdf_with_positions(path):

    words_data = []

    try:
        with pdfplumber.open(path) as pdf:

            for page_num, page in enumerate(pdf.pages, start=1):

                try:
                    words = page.extract_words()

                    for w in words:

                        words_data.append({
                            "text": w.get("text", ""),
                            "x0": float(w.get("x0", 0)),
                            "x1": float(w.get("x1", 0)),
                            "top": float(w.get("top", 0)),
                            "bottom": float(w.get("bottom", 0)),
                            "page": page_num
                        })

                except Exception as e:
                    print(f"⚠️ Word extraction error on page {page_num}:", e)
                    continue

    except Exception as e:
        print("❌ PDF position parsing failed:", e)
        return []

    print("PDF WORDS EXTRACTED:", len(words_data))

    return words_data


# -----------------------------
# 📄 DOCX PARSER
# -----------------------------
def parse_docx(path):

    text = []

    try:
        doc = docx.Document(path)

        # ✅ Paragraphs
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                text.append(p.text.strip())

        # ✅ Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text.append(cell.text.strip())

    except Exception as e:
        print("❌ DOCX parsing failed:", e)
        return ""

    full_text = "\n".join(text)

    print("DOCX TEXT LENGTH:", len(full_text))
    print("DOCX SAMPLE:", full_text[:300])

    return full_text.strip()


# -----------------------------
# 📄 TXT PARSER
# -----------------------------
def parse_txt(path):

    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    except Exception as e:
        print("❌ TXT parsing failed:", e)
        return ""

    print("TXT TEXT LENGTH:", len(text))

    return text.strip()